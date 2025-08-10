import serial
import RPi.GPIO as GPIO
from mfrc522 import SimpleMFRC522
import time
import base64
import os
import threading
import speech_recognition as sr
from queue import Queue, Empty  
import pyttsx3
from google import genai
from google.genai import types
from docx import Document
from datetime import datetime
import requests
import json
from dotenv import load_dotenv

load_dotenv() 

# GPIO Setup
button_pin = 17  # Change to your GPIO pin
GPIO.setmode(GPIO.BCM)
GPIO.setup(button_pin, GPIO.IN, pull_up_down=GPIO.PUD_UP)
# Initializing serial connection to Arduino
arduino_serial = serial.Serial('/dev/ttyUSB0', 115200, timeout=1)

ID = None

def read_rfid():
    """
    Function to read patient ID from an RFID card using MFRC522 reader.
    """
    global ID
    reader = SimpleMFRC522()
    
    try:
        read_text("Hold an RFID card near the reader...")
        # Read the RFID card (this will wait until a card is detected)
        id, text = reader.read()
        ID = str(id)  # Convert the ID to string and store it
        read_text("ID scanned successfully")
        return ID
    except Exception as e:
        read_text("Error reading RFID card: {e}")
        return None
        
    finally:
        GPIO.cleanup()

def read_arduino_vitals():
    """Reads temperature, pulse, and SpO2 from Arduino over serial."""
    try:
        line = arduino_serial.readline().decode('utf-8').strip()
        if line:
            data = {}
            parts = line.split(',')
            for part in parts:
                key, value = part.split(':')
                data[key] = value
            return data
    except Exception as e:
        print(f"Error reading Arduino data: {e}")
        return None

def upload_report_to_server(report_text, ID):
    """Uploads the generated report to the Nexus Medical Backend"""

    if not patient_id:
        print("No patient ID available for upload")
        return False
        
    # API configuration
    base_url = "https://nexus-medi-backend.onrender.com"
    endpoint = "/api/v1/reports"
    patient_id = ID
    auth_key = "nexusrobogenn0825"

    # Prepare the request
    url = f"{base_url}{endpoint}"
    headers = {
        'Content-Type': 'application/json',
        'Authorization': auth_key
    }
    payload = {
        "patient_id": patient_id,
        "report_summary": report_text,
        "created_by": "Robot Nexus",
        "isconfidential": True,
        "status": "Not Responded"
    }

    try:
        print("\nUploading report to server...")
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=10)
        response.raise_for_status()

        result = response.json()
        print("\n✅ Report uploaded successfully!")
        print(f"Report ID: {result['data']['report_id']}")
        print(f"Status: {result['data']['status']}")
        return True

    except requests.exceptions.RequestException as e:
        print("\n❌ Error uploading report:")
        if hasattr(e, 'response') and e.response:
            try:
                error_data = e.response.json()
                print(f"Status Code: {e.response.status_code}")
                print(f"Error: {error_data.get('error', 'Unknown error')}")
            except ValueError:
                print(f"Status Code: {e.response.status_code}")
                print(f"Response Text: {e.response.text}")
        else:
            print(f"Connection Error: {str(e)}")
        return False


def save_report_to_word(report_text):
    doc = Document()
    doc.add_heading("Consultation Report", 0)
    
    for line in report_text.split('\n'):
        if line.strip():  # Skip empty lines
            doc.add_paragraph(line.strip())

    # Get the script's directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Build the filename with full path
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"Consultation_Report_{timestamp}.docx"
    full_path = os.path.join(script_dir, filename)
    
    doc.save(full_path)
    print(f"\nReport saved as: {full_path}")

def speech_to_text_continuous(stop_event):
    audio_queue = Queue()
    recognizer = sr.Recognizer()
    combined_text = []  # Store text from each chunk

    def record_audio():
        with sr.Microphone() as source:
            recognizer.adjust_for_ambient_noise(source)
            print("Recording... (Release button to stop)")
            
            while not stop_event.is_set():
                try:
                    audio = recognizer.listen(source, timeout=0.1, phrase_time_limit=10)
                    audio_queue.put(audio)
                except sr.WaitTimeoutError:
                    # Handle timeout between audio chunks
                    continue
                except Exception as e:
                    print(f"Recording error: {e}")
                    break

    def process_audio():
        while not stop_event.is_set() or not audio_queue.empty():
            try:
                audio = audio_queue.get(timeout=1)
                text = recognizer.recognize_google(audio)
                combined_text.append(text)
                print(f"Partial: {text}")
            except Empty:
                continue
            except sr.UnknownValueError:
                pass  # Skip unrecognized audio
            except sr.RequestError as e:
                print("Skipping a segment due to network error.")
                continue

    # Start recording and processing threads
    record_thread = threading.Thread(target=record_audio, daemon=True)
    process_thread = threading.Thread(target=process_audio, daemon=True)
    
    record_thread.start()
    process_thread.start()
    
    # Wait for stop event
    while not stop_event.is_set():
        time.sleep(0.1)
    
    print("\nStopping recording...")
    
    # Wait for threads to finish
    record_thread.join(timeout=2.0)
    process_thread.join(timeout=2.0)
    
    return " ".join(combined_text)

def read_text(text):
    """Reads the given text using a text-to-speech engine."""
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    
    # Try to find a female voice
    female_voice = None
    for voice in voices:
        if "female" in voice.name.lower() or "Zira" in voice.name or "Samantha" in voice.name:
            female_voice = voice.id
            break

    if female_voice:
        engine.setProperty('voice', female_voice)
    else:
        engine.setProperty('voice', voices[0].id)

    engine.setProperty('rate', 150)
    engine.say(text)
    engine.runAndWait()

def generate(user_input, chat_history):
    client = genai.Client(
        api_key=os.getenv("GEMINI_API_KEY"), 
    )

    model = "gemini-2.0-flash"

    contents = []
    contents.append(types.Content(role="user",parts=[types.Part.from_text(text= medical_prompt)]))

    for message in chat_history:
        contents.append(types.Content(role=message["role"], parts=[types.Part.from_text(text=message["text"])]))

    contents.append(types.Content(role="user", parts=[types.Part.from_text(text=user_input)]))

    tools = [
        types.Tool(google_search=types.GoogleSearch())
    ]
    generate_content_config = types.GenerateContentConfig(
        temperature=1,
        top_p=0.95,
        top_k=40,
        max_output_tokens=8192,
        tools=tools,
        response_mime_type="text/plain",
    )

    response = ""
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        print(chunk.text, end="")
        response += chunk.text

    return response

# Load medical prompt
script_dir = os.path.dirname(os.path.abspath(__file__))
prompt_path = os.path.join(script_dir, "medical_prompt.txt")
with open(prompt_path, "r", encoding="utf-8") as file:
    medical_prompt = file.read()

chat_history = []

try:
    read_text("Press and release the button to stop recording...")
    
    while True:
        # Read Arduino vitals
        vitals = read_arduino_vitals()
        # Wait for button press (LOW)
        if not GPIO.input(button_pin):
            Patient_ID = read_rfid()
            if not patient_id:
                    read_text("Failed to read patient ID. Please try again.")
                    continue
            stop_event = threading.Event()
                        
            read_text("recording in progress...please begin consultation session")
            
            # Start recording in a separate thread
            recording_thread = threading.Thread(
                target=lambda: globals().update({'result': speech_to_text_continuous(stop_event)}),
                daemon=True
            )
            recording_thread.start()
            
            # Wait for button release
            while not GPIO.input(button_pin):
                time.sleep(0.1)
            
            # Button released - stop recording
            stop_event.set()
            recording_thread.join()
            
            read_text("Session ended. Generating report...")

            
            # Generate report (include vitals if available)
            if 'result' in globals() and globals()['result']:
                consultation_text = globals()['result']
                if vitals:
                    consultation_text += f"\n\nVitals:\n- Temperature: {vitals.get('TEMP')}°C\n- Pulse: {vitals.get('PULSE')} BPM\n- SpO2: {vitals.get('SpO2')}%"
                
                response = generate(consultation_text, chat_history)

                # Upload to server
                if upload_report_to_server(cleaned_response,ID):
                    read_text("Report has been successfully uploaded to the patient's records.")
                else:
                    read_text("Warning: Could not upload report to server. Local copy has been saved.")

                
                # Update chat history
                chat_history.append({"role": "user", "text": globals()['result']})
                chat_history.append({"role": "model", "text": response})

except KeyboardInterrupt:
    read_text("Exiting...")
finally:
    GPIO.cleanup()
